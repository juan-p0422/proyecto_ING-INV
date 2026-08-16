from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\Juan Pantoja\Desktop\proyecto_ING-INV")
SOURCE = ROOT / "docs" / "31-reporte-final-pdf-ready.md"
LOGO = ROOT / "docs" / "assets" / "ceti-logo.jpg"
OUTPUT = ROOT / "evidence" / "final-pdf" / "II_GLOBAL_23110022_8C.docx"

BLUE = "0B67A8"
DARK_BLUE = "173B67"
ORANGE = "F7901E"
INK = "172033"
MUTED = "526175"
LIGHT_BLUE = "EAF3FA"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = OxmlElement("w:r")
    r.append(begin)
    r.append(instr)
    r.append(separate)
    r.append(text)
    r.append(end)
    paragraph._p.append(r)


def set_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_picture_alt(inline_shape, alt_text: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text[:80])


def fit_image(path: Path, max_width_in: float, max_height_in: float) -> tuple[float, float]:
    with Image.open(path) as image:
        width, height = image.size
    ratio = width / height if height else 1
    out_w = max_width_in
    out_h = out_w / ratio
    if out_h > max_height_in:
        out_h = max_height_in
        out_w = out_h * ratio
    return out_w, out_h


def add_picture(paragraph, path: Path, max_width_in: float, max_height_in: float, alt: str) -> None:
    width, height = fit_image(path, max_width_in, max_height_in)
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width), height=Inches(height))
    set_picture_alt(shape, alt)


def strip_markup(text: str) -> str:
    text = re.sub(r"!\[([^]]*)\]\([^)]+\)(?:\{[^}]*\})?", r"\1", text)
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "")
    text = re.sub(r"<[^>]+>", "", text)
    return text.strip()


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^]]+\]\([^)]+\)|https?://\S+)")


def add_inline(paragraph, text: str, base_size=11, base_color=INK) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=base_size, color=base_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=max(8.5, base_size - 1), color=DARK_BLUE)
            run.font.highlight_color = None
        elif token.startswith("["):
            m = re.match(r"\[([^]]+)\]\(([^)]+)\)", token)
            label, target = m.group(1), m.group(2)
            run = paragraph.add_run(f"{label} ({target})")
            set_font(run, size=base_size, color=BLUE)
            run.underline = True
        else:
            run = paragraph.add_run(token.rstrip(".,;"))
            set_font(run, size=base_size, color=BLUE)
            run.underline = True
            if token[-1:] in ".,;":
                trail = paragraph.add_run(token[-1])
                set_font(trail, size=base_size, color=base_color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size=base_size, color=base_color)


def add_paragraph(doc, text: str, style=None, align=None, italic=False, color=INK, size=11):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    add_inline(paragraph, text, base_size=size, base_color=color)
    if italic:
        for run in paragraph.runs:
            run.italic = True
    return paragraph


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def column_widths(count: int) -> list[int]:
    if count == 1:
        return [TABLE_WIDTH_DXA]
    if count == 2:
        return [2800, 6560]
    if count == 3:
        return [2100, 3760, 3500]
    if count == 4:
        return [1700, 2700, 2300, 2660]
    if count == 5:
        return [1250, 2050, 1700, 2500, 1860]
    base = TABLE_WIDTH_DXA // count
    widths = [base] * count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def add_data_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    widths = column_widths(columns)
    for r_idx, row in enumerate(rows):
        for c_idx in range(columns):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.05
            value = row[c_idx] if c_idx < len(row) else ""
            add_inline(paragraph, value, base_size=8.6 if columns >= 4 else 9.2)
            if r_idx == 0:
                set_cell_shading(cell, DARK_BLUE)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
        prevent_row_split(table.rows[r_idx])
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_callout(doc, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    add_inline(paragraph, text, base_size=10.2, base_color=DARK_BLUE)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines: Iterable[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_font(run, name="Consolas", size=8.7, color=INK)
        if idx < len(list(lines)) - 1:
            run.add_break()
    set_table_geometry(table, [TABLE_WIDTH_DXA])


FIGURE_RE = re.compile(r'<figure><img src="([^"]+)"><figcaption>(.*?)</figcaption></figure>')


def add_figure_grid(doc, entries: list[tuple[str, str]]) -> None:
    for start in range(0, len(entries), 2):
        pair = entries[start:start + 2]
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, cell in enumerate(table.rows[0].cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, top=45, start=70, bottom=45, end=70)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if idx < len(pair):
                src, caption_html = pair[idx]
                caption = strip_markup(caption_html)
                image_path = (ROOT / "docs" / src).resolve()
                # Keep evidence legible while ensuring three or four grid rows stay
                # inside the printable area instead of displacing headers/footers.
                add_picture(paragraph, image_path, 2.65, 1.55, caption)
                cp = cell.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(3)
                cp.paragraph_format.space_after = Pt(0)
                add_inline(cp, caption, base_size=8.0, base_color=MUTED)
                for run in cp.runs:
                    run.italic = True
            else:
                paragraph.add_run("")
        # Word may place an unsplittable image row behind the page header when it
        # is carried to the next page. The row is compact enough to paginate
        # naturally, so leave splitting enabled for stable print layout.
        set_table_geometry(table, [4680, 4680], indent_dxa=0)
        table._tbl.tblPr.append(OxmlElement("w:tblCellMar"))
        doc.add_paragraph().paragraph_format.space_after = Pt(1)
        pair_number = start // 2 + 1
        if pair_number % 3 == 0 and start + 2 < len(entries):
            doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True
    # Define explicit odd/even headers and footers. Microsoft Word can otherwise
    # reuse stale field geometry while exporting long, image-heavy documents,
    # which caused even pages to lose the header and show only the last digit.
    doc.settings.odd_and_even_pages_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Heading 1", 16, DARK_BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    doc.core_properties.title = "EduRoom - Reporte final de práctica"
    doc.core_properties.subject = "Ingeniería inversa ética de Google Classroom y réplica académica independiente"
    doc.core_properties.author = "Juan Oswaldo Emilio Olivares Pantoja"
    doc.core_properties.keywords = "EduRoom, Ingeniería Inversa, Google Classroom, CETI, 8C"


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    add_picture(p, LOGO, 2.15, 2.15, "Logo institucional del Centro de Enseñanza Técnica Industrial")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("CENTRO DE ENSEÑANZA TÉCNICA INDUSTRIAL")
    set_font(run, size=12, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("EDUROOM")
    set_font(run, size=29, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Reporte final de práctica")
    set_font(run, size=17, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Ingeniería inversa ética de Google Classroom\ny réplica académica independiente")
    set_font(run, size=12.5, color=MUTED, italic=True)

    metadata = [
        ("Materia", "Ingeniería Inversa"),
        ("Aplicación objetivo", "Google Classroom"),
        ("Alumno", "Juan Oswaldo Emilio Olivares Pantoja"),
        ("Número de registro", "23110022"),
        ("Grupo", "8C"),
        ("Fecha", "16 de agosto de 2026"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}: ")
        set_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(value)
        set_font(r2, size=10.5, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    add_inline(p, "Repositorio: https://github.com/juan-p0422/proyecto_ING-INV", base_size=8.8, base_color=MUTED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, "Deploy: https://eduroom-znb0.onrender.com", base_size=8.8, base_color=MUTED)
    doc.add_page_break()


def add_static_toc(doc: Document) -> None:
    p = doc.add_paragraph("Índice", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    sections = [
        "Nota de autorización y protección de secretos",
        "Lista de figuras y evidencias",
        "1. Introducción",
        "2. Objetivos",
        "3. Marco teórico",
        "4. Aplicación analizada: Google Classroom",
        "5. Ingeniería inversa y metodología ética",
        "6. Análisis dinámico, Performance y Memory",
        "7. Análisis de vulnerabilidades",
        "8. Reconstrucción de estructuras",
        "9. Réplica EduRoom",
        "10. Cifrado y ofuscación",
        "11. Checksum e integridad",
        "12. Antireversing y limitaciones del cliente web",
        "13. Despliegue",
        "14. Pruebas y endpoints",
        "15. Evidencias y cumplimiento",
        "16. Conclusiones",
        "17. Integración documental",
        "Anexo visual de evidencias existentes",
    ]
    for item in sections:
        paragraph = doc.add_paragraph(style="List Number")
        add_inline(paragraph, item, base_size=10.2)
    doc.add_page_break()


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_static_toc(doc)

    start = next(i for i, line in enumerate(lines) if line.strip() == "# Nota de autorización y protección de secretos")
    i = start
    in_code = False
    code_lines: list[str] = []
    skip_index = False
    skip_export = False
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if line == "# Índice":
            skip_index = True
            i += 1
            continue
        if skip_index:
            if line == "# Lista de figuras y evidencias":
                skip_index = False
            else:
                i += 1
                continue

        if line == "# 18. Instrucciones de exportación":
            skip_export = True
            i += 1
            continue
        if skip_export:
            if line == "# 19. Anexo visual de evidencias existentes":
                skip_export = False
                doc.add_page_break()
                doc.add_paragraph("Anexo visual de evidencias existentes", style="Heading 1")
                i += 1
                continue
            i += 1
            continue

        if line.startswith("<style"):
            while i < len(lines) and "</style>" not in lines[i]:
                i += 1
            i += 1
            continue
        if line.startswith("<div class=\"page-break"):
            i += 1
            continue
        if line == '<div class="figure-grid">':
            entries: list[tuple[str, str]] = []
            i += 1
            while i < len(lines) and lines[i].strip() != "</div>":
                match = FIGURE_RE.search(lines[i].strip())
                if match:
                    entries.append((match.group(1), match.group(2)))
                i += 1
            add_figure_grid(doc, entries)
            i += 1
            continue
        if line.startswith("<") and line.endswith(">"):
            i += 1
            continue
        if line == "---":
            i += 1
            continue
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                table = doc.add_table(rows=1, cols=1)
                cell = table.cell(0, 0)
                set_cell_shading(cell, LIGHT_GRAY)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                for idx, code_line in enumerate(code_lines):
                    run = paragraph.add_run(code_line)
                    set_font(run, name="Consolas", size=8.5, color=INK)
                    if idx < len(code_lines) - 1:
                        run.add_break()
                set_table_geometry(table, [TABLE_WIDTH_DXA])
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_data_table(doc, rows)
            continue
        if line == "# 17. Integración documental":
            doc.add_page_break()
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif line == "## 19.5 Categorías técnicas sin imagen disponible":
            break
        elif line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif line in (
            "## 19.2 Google Classroom - evidencia dinámica",
            "## 19.3 EduRoom - serie principal de la réplica",
            "## 19.4 EduRoom - serie complementaria",
        ):
            doc.add_page_break()
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
        elif line.startswith("> "):
            add_callout(doc, line[2:].strip())
        elif re.match(r"^- \[[ xX]\] ", line):
            checked = line[3].lower() == "x"
            text = re.sub(r"^- \[[ xX]\] ", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, ("Completado: " if checked else "Pendiente: ") + text, base_size=10.2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:], base_size=10.5)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line), base_size=10.5)
        elif line.startswith("!["):
            match = re.match(r"!\[([^]]*)\]\(([^)]+)\)(?:\{[^}]*\})?", line)
            if match:
                alt, src = match.group(1), match.group(2)
                path = (SOURCE.parent / src).resolve()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_picture(p, path, 5.9, 4.8, alt)
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(cp, alt, base_size=9.0, base_color=MUTED)
                for run in cp.runs:
                    run.italic = True
        else:
            add_paragraph(doc, line)
        i += 1

    doc.save(OUTPUT)
    print(f"DOCX creado: {OUTPUT}")


if __name__ == "__main__":
    build()
